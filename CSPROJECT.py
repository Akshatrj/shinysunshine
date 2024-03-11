from datetime import datetime
import csv
import subprocess
from PIL import Image
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import mysql.connector as sqltor
import os
os.environ['PYGAME_HIDE_SUPPORT_PROMPT'] = "hide"
import pygame


mycon= sqltor.connect(host="localhost",user="root",password="12345",database="store")
mycursor= mycon.cursor()

custom_width = 4.25 * 72  # Adjusted for better readability(to convert cm to dpi)
custom_length = 11 * 72  # Letter-sized page
filename = f"bill_{datetime.now().strftime('%Y%m%d%H%M%S')}.pdf"#converts datetime to string
c = canvas.Canvas(filename, pagesize=(custom_width, custom_length))
f=open("Customer_and_Sales_Data.csv", "a",newline="")
m=csv.writer(f)

def check_expired_products():
    today = datetime.now().date()#get date time stored in a variable 
    query = "SELECT product_name, expiry_date FROM inventory WHERE expiry_date < %s"
    mycursor.execute(query, (today,))#check the expired by comparing date of today and date set in database
    expired_products = mycursor.fetchall()#fetching data

    if expired_products:
        print("Expired Products:")
        for product in expired_products:
            print(f"{product[0]} - Expired on {product[1]}")
    else:#print exp product
        print("No expired products.")


def handle_input_int(prompt):#had made a fuction to deal with error
    while True:
        try:
            value = int(input(prompt)) 
            return value
        except ValueError:
            print("Invalid input. Please enter a valid number.")


def generate_bill():#to generate bills
    f = open("Customer_and_Sales_Data.csv", "a", newline='')
    m = csv.writer(f)
    today = datetime.now().date()#get date time stored in a variable 

    try:
        customer_Name = input("Enter Customer Name: ")
        Customer_PhoneNumber = handle_input_int("Enter Customer Number: ")

        if len(str(Customer_PhoneNumber)) == 10:
            items = []
            choice = "Y"

            while choice.upper() == "Y":
                product_code = input("Enter the product code: ")
                quantity1 = handle_input_int("Enter the quantity: ")

                query = "SELECT product_name, selling_price, quantity,expiry_date FROM inventory WHERE product_code = %s"
                mycursor.execute(query, (product_code,))
                product_data = mycursor.fetchone()

                if product_data:
                    product_name, selling_price, quantity,expiry_date = product_data

                    if quantity1 <= quantity and expiry_date > today :
                        total_price = selling_price * quantity1

                        items.append({
                            "product_name": product_name,
                            "quantity": quantity1,
                            "selling_price": selling_price,
                            "total_price": total_price
                        })
                    else:
                        print("Product out of stock or Expired")

                else:
                    print("Product not found!!")

                choice = input("Enter 'Y' to add more products: ")

            if items:
                doc = Document()

                # Add header
                doc.add_heading("Shopper's Grocery", level=1)
                doc.add_paragraph("=" * 50)
                doc.add_heading("Receipt")
                doc.add_paragraph(f"Customer Name:{customer_Name}")
                doc.add_paragraph(f"Customer No.-{Customer_PhoneNumber}")
                doc.add_paragraph("=" * 50)
                for item in items:
                    # Add item details
                    doc.add_paragraph(f"Product: {item['product_name']}")
                    doc.add_paragraph(f"Quantity: {item['quantity']}")
                    doc.add_paragraph(f"Selling Price per unit: Rs{item['selling_price']}")
                    doc.add_paragraph(f"Total Price: Rs{item['total_price']}")
                    doc.add_paragraph("=" * 20)

                total_bill = sum(item['total_price'] for item in items)

                # Add total bill
                doc.add_paragraph(f"Total Bill: Rs{total_bill}")

                # Add store information
                doc.add_paragraph("=" * 50)
                doc.add_heading("Store Information:", level=2)
                doc.add_paragraph("Address: Fatehgarh, Farrukhabad 209601")
                doc.add_paragraph("Contact Us at 9090909090")
                doc.add_paragraph("Email: shoppergrocery@gmail.com")

                # Add footer
                doc.add_paragraph("=" * 50)
                doc.add_paragraph("Thank You Please visit Us Again!!")
                doc.add_paragraph("☺ ☺ ☺ ☺ ☺ ☺")

                # Save the document
                filename = f"bill_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"

                # print("Please Select the Payment Method")
                print("1. Cash")
                print("2. Online Banking")
                p = input("What's Your Payment Method Please press(1/2) to enter:")
                if p == "1":
                    payment_confirmation = input("Confirm payment (yes/no): ").lower()

                    if payment_confirmation == "yes":
                        print("Payment successful. Thank You for shopping!")
                        print("Please Visit Us Again......")
                        doc.save(filename)
                        print("Bill generated successfully.")
                        pygame.init()
                        file_path = "D:\\Computer Project\\Voice.mp3"
                        mixer = pygame.mixer
                        sound = mixer.Sound(file_path)
                        sound.play()
                        while mixer.get_busy():
                            pygame.time.wait(100)
                        pygame.quit()

                        subprocess.Popen(["start", filename], shell=True)

                        for item in items:
                            update_query = "UPDATE inventory SET quantity = quantity - %s WHERE product_code = %s"
                            mycursor.execute(update_query, (item['quantity'], product_code))
                            mycon.commit()

                            l = [customer_Name, Customer_PhoneNumber, product_code, item['quantity'],item['product_name'],item['selling_price'], item['total_price']]
                            m.writerow(l)
                    else:
                            print("Payment not confirmed. No changes to inventory.")
                elif p == "2":
                    image_path = r"D:\Computer Project\qr.PNG"
                    subprocess.Popen(["start", "", image_path], shell=True)


                    payment_confirmation = input("Confirm payment (yes/no): ").lower()

                    if payment_confirmation == "yes":
                        print("Payment successful. Thank You for shopping!")
                        print("Please Visit Us Again......")
                        doc.save(filename)
                        print("Bill generated successfully.")
                        pygame.init()
                        file_path = "D:\\Computer Project\\Voice.mp3"
                        mixer = pygame.mixer
                        sound = mixer.Sound(file_path)
                        sound.play()
                        while mixer.get_busy():
                            pygame.time.wait(100)
                        pygame.quit()

                        subprocess.Popen(["start", filename], shell=True)

                        for item in items:
                            update_query = "UPDATE inventory SET quantity = quantity - %s WHERE product_code = %s"
                            mycursor.execute(update_query, (item['quantity'], product_code))
                            mycon.commit()

                            l = [customer_Name, Customer_PhoneNumber, product_code, item['quantity'],item['product_name'],item['selling_price'], item['total_price']]
                            m.writerow(l)
                    else:
                            print("Payment not confirmed. No changes to inventory.")
                else:
                    print("Please Enter the Valid Choice")

            else:
                print("No items in the bill. Please add items.")

        else:
            print("Invalid Phone Number!!")

    except Exception as e:
        print(f"An error occurred: {e}")
    f.close()


def add_product():
    ch = "y"
    while ch in ["Y", "y"]:
        product_name = input("Enter the product name: ")
        product_code = input("Enter the product code: ")
        cost_price = float(input("Enter the cost price per unit: "))
        selling_price = float(input("Enter the selling price per unit: "))
        quantity = float(input("Enter the initial quantity: "))
        expiry_date_str = input("Enter the expiry date (YYYY-MM-DD): ")
        expiry_date = datetime.strptime(expiry_date_str, "%Y-%m-%d").date()

        query = "INSERT INTO inventory (product_name, product_code, cost_price, selling_price, quantity, expiry_date) VALUES (%s, %s, %s, %s, %s, %s)"
        mycursor.execute(query, (product_name, product_code, cost_price, selling_price, quantity, expiry_date))
        mycon.commit()
        print(f"{quantity} units of {product_name} have been added to the inventory.")

        ch = input("For more entries press Y otherwise press any other letter.")


def check_profit_loss():
    query = "SELECT product_name, quantity, cost_price, selling_price, expiry_date FROM inventory"
    mycursor.execute(query)
    products_data = mycursor.fetchall()

    total_profit = 0
    loss = 0

    print("Profit and Loss Report")
    print('=' * 94)
    print(f"{'Product': <15}{'Quantity': <10}{'Unit Price':<15}{'Cost Price': <15}{'Selling Price': <15}{'Profit': <15}{'Loss':<15}")
    print("=" * 94)

    for product_data in products_data:
        product_name, quantity, cost_price, selling_price, expiry_date = product_data
        cp=cost_price
        total_cost_price = cost_price * quantity
        total_selling_price = selling_price * quantity
        p=0
        l = 0
        today = datetime.now().date()
        if expiry_date < today:
            a = cost_price * quantity
            l += a
        else:
            profit = total_selling_price - total_cost_price
            p+=profit
        loss += l
        total_profit += p

        print(f"{product_name: <15}{quantity: <10}{cp:<15.2f}{total_cost_price: <15.2f}{total_selling_price: <15.2f}{p: <15.2f}{l:<15.2f}")

    print("=" * 94)
    print("Total Profit: Rs",total_profit)
    print("Total Loss Rs", loss)


def view_sales():
    f=open("Customer_and_Sales_Data.csv", "r",)
    r=csv.reader(f)
    for i in r:
        print(i)
    f.close()


def admin_menu():
    while True:
        print("\nAdmin Menu:")
        print("1. Check Expired Products")
        print("2. Generate Bill")
        print("3. Check Profit and Loss")
        print("4. Add Product to Inventory")
        print("5. View Sales And Customer Data")
        print("6. Exit")

        choice = input("Enter your choice (1/2/3/4/5/6): ")

        if choice == "1":
            check_expired_products()
        elif choice == "2":
            generate_bill()
        elif choice == "3":
            check_profit_loss()
        elif choice == "4":
            add_product()
        elif choice == "5":
            view_sales()
        elif choice == "6":
            break
        else:
            print("Invalid choice. Please enter a valid option.")


def user_menu():
    while True:
        print("\nUser Menu:")
        print("1. Check Expired Products")
        print("2. Generate Bill")
        print("3. Exit")

        choice = input("Enter your choice (1/2/3): ")

        if choice == "1":
            check_expired_products()
        elif choice == "2":
            generate_bill()
        elif choice == "3":
            
            break
        else:
            print("Invalid choice. Please enter a valid option.")


def login():
    username = input("Enter your username: ")
    password = input("Enter your password: ")

    admin_credentials = {"admin": "admin123"}
    user_credentials = {"user": "user123"}

    if username in admin_credentials and password == admin_credentials[username]:
        print("Welcome, Admin!")
        admin_menu()
    elif username in user_credentials and password == user_credentials[username]:
        print("Welcome, User!")
        user_menu()
    else:
        print("Invalid credentials. Please try again.")


def main():
    login()


if __name__ == "__main__":
    main()

mycursor.close()
mycon.close()

